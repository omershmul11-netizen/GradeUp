#!/usr/bin/env python3
"""Render validated Bagrut 471 JSON specifications as exam-style PDFs."""

import ast
import json
import math
import os
from pathlib import Path
import re
import shutil
import subprocess
import sys
import tempfile

import numpy as np
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.shared import Cm, Pt
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


APP_ROOT = Path(__file__).resolve().parents[1]
FONT_DIR = APP_ROOT / "bagrut471" / "assets" / "fonts"
HEBREW_REGULAR = "JSTVKF+FbDavidNewPro-Regular"
HEBREW_BOLD = "JSTVKF+FbDavidNewPro-Bold"
MATH_FONT = "Times New Roman"


def executable(name):
    value = shutil.which(name)
    if not value:
        raise RuntimeError(f"Required executable is unavailable: {name}")
    return value


def ltr(text):
    return "\u202a" + str(text) + "\u202c"


def set_paragraph_rtl(paragraph):
    props = paragraph._p.get_or_add_pPr()
    if props.find(qn("w:bidi")) is None:
        props.append(OxmlElement("w:bidi"))
    # LibreOffice uses physical-left as the logical start for this bidi setup.
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def set_run_font(run, size=12, bold=False, italic=False, font_name=None):
    face = font_name or (HEBREW_BOLD if bold else HEBREW_REGULAR)
    run.font.name = face
    run_props = run._element.get_or_add_rPr()
    fonts = run_props.get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "cs"):
        fonts.set(qn(f"w:{key}"), face)
    run.font.size = Pt(size)
    run.font.bold = bool(bold and font_name is not None)
    run.font.italic = italic
    language = OxmlElement("w:lang")
    language.set(qn("w:val"), "he-IL")
    language.set(qn("w:bidi"), "he-IL")
    run_props.append(language)


def add_rtl_paragraph(doc, text="", size=12, bold=False, after=3, before=0,
                      line=1.08):
    paragraph = doc.add_paragraph()
    set_paragraph_rtl(paragraph)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.line_spacing = line
    if text:
        set_run_font(paragraph.add_run(text), size=size, bold=bold)
    return paragraph


def add_math_text(paragraph, text, size=11, bold=False, italic=True):
    run = paragraph.add_run(ltr(text))
    set_run_font(run, size=size, bold=bold, italic=italic, font_name=MATH_FONT)
    rtl_node = OxmlElement("w:rtl")
    rtl_node.set(qn("w:val"), "0")
    run._element.get_or_add_rPr().append(rtl_node)
    return run


def append_mixed(paragraph, parts, size=12, bold=False):
    for kind, value in normalize_parts(parts):
        if kind == "he":
            set_run_font(paragraph.add_run(value), size=size, bold=bold)
        else:
            add_math_text(paragraph, value, size=max(9, size - 1), bold=bold)


def set_bottom_border(paragraph, color="D5DEE8", size="6", space="4"):
    props = paragraph._p.get_or_add_pPr()
    borders = props.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        props.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def set_cell_bottom_border(cell, color="DCE4EA", size="3"):
    props = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "nil")
        borders.append(border)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    props.append(borders)


def add_worksheet_header(doc, question):
    brand = add_rtl_paragraph(doc, after=2)
    brand_run = brand.add_run("GradeUp  |  דף עבודה במתמטיקה")
    set_run_font(brand_run, size=9, bold=True)
    brand_run.font.color.rgb = RGBColor(31, 79, 112)
    set_bottom_border(brand, color="9FC3D5", size="8", space="5")

    title = add_rtl_paragraph(doc, str(question.get("title", "דף עבודה")),
                              size=18, bold=True, after=2, before=5)
    title.paragraph_format.keep_with_next = True
    meta = add_rtl_paragraph(doc, after=5)
    set_run_font(meta.add_run(str(question.get("topic_label", ""))), size=10, bold=True)
    set_run_font(meta.add_run("  |  "), size=10)
    set_run_font(meta.add_run(str(question.get("unit_label", "י״א | 4 יח״ל"))), size=10)
    set_run_font(meta.add_run("  |  "), size=10)
    set_run_font(meta.add_run(str(question.get("difficulty_label", "מדורג"))), size=10)

    identity = add_rtl_paragraph(
        doc,
        "שם: ____________________    כיתה: __________    תאריך: __________",
        size=10, after=6,
    )
    set_bottom_border(identity, color="D5DEE8", size="4", space="5")
    instructions = str(question.get("instructions", "פתרו לפי הסדר והציגו דרך."))
    info = add_rtl_paragraph(doc, "דרך עבודה: " + instructions, size=10, after=7)
    info.paragraph_format.keep_with_next = True


def add_section_heading(doc, title, instruction=""):
    paragraph = add_rtl_paragraph(doc, after=3, before=5)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(str(title))
    set_run_font(run, size=12, bold=True)
    run.font.color.rgb = RGBColor(31, 79, 112)
    if instruction:
        set_run_font(paragraph.add_run("  |  " + str(instruction)), size=9)
    set_bottom_border(paragraph, color="9FC3D5", size="6", space="3")


def add_worksheet_exercise(doc, exercise):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(15.9)
    table.columns[1].width = Cm(1.1)
    table_props = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "nil")
        borders.append(border)
    table_props.append(borders)

    # LibreOffice lays out the first table cell on the right in this RTL document.
    # Put the content there and reserve the second (physical-left) cell for numbering.
    content_cell, number_cell = table.rows[0].cells
    number_cell.width = Cm(1.1)
    content_cell.width = Cm(15.9)
    number_paragraph = number_cell.paragraphs[0]
    number_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    number_paragraph.paragraph_format.keep_with_next = True
    number_run = number_paragraph.add_run(f"{int(exercise.get('number'))}.")
    set_run_font(number_run, size=11, bold=True, font_name=MATH_FONT)

    paragraph = content_cell.paragraphs[0]
    set_paragraph_rtl(paragraph)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.keep_with_next = True
    append_mixed(paragraph, exercise.get("parts", []), size=11)
    if exercise.get("formula"):
        set_run_font(paragraph.add_run(" "), size=11)
        equation = OxmlElement("m:oMath")
        append_formula(equation, exercise.get("formula"))
        paragraph._p.append(equation)

    line_count = min(max(int(exercise.get("workspace_lines", 2)), 1), 5)
    workspace_table = doc.add_table(rows=1, cols=1)
    workspace_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    workspace_table.autofit = False
    workspace_table.columns[0].width = Cm(17.0)
    workspace_row = workspace_table.rows[0]
    workspace_row.height = Cm(0.55 * line_count)
    workspace_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    workspace_cell = workspace_row.cells[0]
    workspace_cell.width = Cm(17.0)
    workspace_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
    set_cell_bottom_border(workspace_cell)
    workspace_paragraph = workspace_cell.paragraphs[0]
    workspace_paragraph.paragraph_format.space_after = Pt(0)
    workspace_paragraph.paragraph_format.space_before = Pt(0)
    workspace_marker = workspace_paragraph.add_run("·")
    set_run_font(workspace_marker, size=8)
    workspace_marker.font.color.rgb = RGBColor(255, 255, 255)


def render_practice_worksheet_docx(question, docx_path):
    doc = new_document()
    add_worksheet_header(doc, question)
    sections = question.get("sections", [])
    if len(sections) != 4:
        raise ValueError("A practice worksheet must contain exactly four sections")
    exercise_count = 0
    for section_index, section in enumerate(sections):
        add_section_heading(doc, section.get("title", ""), section.get("instruction", ""))
        exercises = section.get("exercises", [])
        if not exercises:
            raise ValueError("A worksheet section cannot be empty")
        for exercise in exercises:
            exercise_count += 1
            if int(exercise.get("number", 0)) != exercise_count:
                raise ValueError("Exercise numbering must be consecutive")
            add_worksheet_exercise(doc, exercise)
        if section_index == 1:
            doc.add_page_break()
    if not 8 <= exercise_count <= 12:
        raise ValueError("A practice worksheet must contain 8-12 exercises")
    doc.save(docx_path)


def normalize_parts(parts):
    if not isinstance(parts, list):
        raise ValueError("parts must be an array")
    clean = []
    for part in parts:
        if not isinstance(part, dict):
            raise ValueError("each part must be an object")
        kind = part.get("kind")
        text = str(part.get("text", ""))
        if kind not in {"he", "math"} or not text:
            raise ValueError("invalid mixed-text part")
        clean.append((kind, text))
    return clean


def add_mixed(doc, parts, size=12, after=3, before=0, bold=False):
    paragraph = add_rtl_paragraph(doc, after=after, before=before)
    for kind, value in normalize_parts(parts):
        if kind == "he":
            set_run_font(paragraph.add_run(value), size=size, bold=bold)
        else:
            add_math_text(paragraph, value, size=max(9, size - 1), bold=bold)
    return paragraph


def add_item(doc, label, parts, size=12, after=3, before=0):
    paragraph = add_rtl_paragraph(doc, after=after, before=before)
    main_sub = re.fullmatch(r"([א-ת]\.)\s*\((\d+)\)", label)
    sub_only = re.fullmatch(r"\((\d+)\)", label)
    if main_sub:
        set_run_font(paragraph.add_run(main_sub.group(1) + "  "), size=size, bold=True)
        add_math_text(paragraph, f"({main_sub.group(2)})  ", size=11,
                      bold=True, italic=False)
    elif sub_only:
        set_run_font(paragraph.add_run("\u2003"), size=size)
        add_math_text(paragraph, f"({sub_only.group(1)})  ", size=11,
                      bold=True, italic=False)
    elif re.fullmatch(r"[א-ת]\.", label):
        set_run_font(paragraph.add_run(label + "  "), size=size, bold=True)
    else:
        raise ValueError(f"Invalid item label: {label}")
    for kind, value in normalize_parts(parts):
        if kind == "he":
            set_run_font(paragraph.add_run(value), size=size)
        else:
            add_math_text(paragraph, value, size=max(9, size - 1))
    return paragraph


def math_run(text, italic=True):
    node = OxmlElement("m:r")
    if italic:
        props = OxmlElement("m:rPr")
        style = OxmlElement("m:sty")
        style.set(qn("m:val"), "i")
        props.append(style)
        node.append(props)
    value = OxmlElement("m:t")
    value.text = str(text)
    node.append(value)
    return node


def formula_node(spec):
    if not isinstance(spec, dict):
        raise ValueError("formula node must be an object")
    node_type = spec.get("type")
    if node_type == "text":
        return math_run(spec.get("value", ""), spec.get("italic", True))
    if node_type == "row":
        return [formula_node(child) for child in spec.get("children", [])]
    if node_type == "sup":
        node = OxmlElement("m:sSup")
        expression = OxmlElement("m:e")
        append_formula(expression, spec.get("base"))
        exponent = OxmlElement("m:sup")
        append_formula(exponent, spec.get("exponent"))
        node.extend([expression, exponent])
        return node
    if node_type == "frac":
        node = OxmlElement("m:f")
        numerator = OxmlElement("m:num")
        denominator = OxmlElement("m:den")
        append_formula(numerator, spec.get("numerator"))
        append_formula(denominator, spec.get("denominator"))
        node.extend([numerator, denominator])
        return node
    if node_type == "sqrt":
        node = OxmlElement("m:rad")
        props = OxmlElement("m:radPr")
        hidden = OxmlElement("m:degHide")
        hidden.set(qn("m:val"), "1")
        props.append(hidden)
        degree = OxmlElement("m:deg")
        expression = OxmlElement("m:e")
        append_formula(expression, spec.get("value"))
        node.extend([props, degree, expression])
        return node
    raise ValueError(f"Unsupported formula node: {node_type}")


def append_formula(parent, spec):
    result = formula_node(spec)
    if isinstance(result, list):
        for child in result:
            parent.append(child)
    else:
        parent.append(result)


def add_intro(doc, question_number, formula):
    paragraph = add_rtl_paragraph(doc, after=4)
    add_math_text(paragraph, f"{int(question_number)}.", size=11,
                  bold=True, italic=False)
    set_run_font(paragraph.add_run("\u2002"), size=12)
    set_run_font(paragraph.add_run("נתונה הפונקציה  "), size=12)
    equation = OxmlElement("m:oMath")
    append_formula(equation, formula)
    paragraph._p.append(equation)


def add_general_intro(doc, question_number, parts):
    paragraph = add_rtl_paragraph(doc, after=5)
    add_math_text(paragraph, f"{int(question_number)}.  ", size=11,
                  bold=True, italic=False)
    for kind, value in normalize_parts(parts):
        if kind == "he":
            set_run_font(paragraph.add_run(value), size=12)
        else:
            add_math_text(paragraph, value, size=11)


def add_data_table(doc, headers, rows, after=5):
    if not isinstance(headers, list) or not 2 <= len(headers) <= 8:
        raise ValueError("data_table requires 2-8 headers")
    if not isinstance(rows, list) or not 1 <= len(rows) <= 12:
        raise ValueError("data_table requires 1-12 rows")
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    values = [headers] + rows
    for row_index, source_row in enumerate(values):
        if not isinstance(source_row, list) or len(source_row) != len(headers):
            raise ValueError("all data_table rows must match the header count")
        target_cells = table.rows[row_index].cells if row_index == 0 else table.add_row().cells
        for cell, value in zip(target_cells, source_row):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            set_paragraph_rtl(paragraph)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(paragraph.add_run(str(value)), size=10, bold=row_index == 0)
    spacer = add_rtl_paragraph(doc, after=after)
    spacer.paragraph_format.space_after = Pt(after)


def add_page_number(section):
    paragraph = section.footer.paragraphs[0]
    for child in list(paragraph._p):
        paragraph._p.remove(child)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])
    set_run_font(run, size=10)


def new_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)
    normal = doc.styles["Normal"]
    normal.font.name = HEBREW_REGULAR
    normal.font.size = Pt(12)
    fonts = normal._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "cs"):
        fonts.set(qn(f"w:{key}"), HEBREW_REGULAR)
    settings = doc.settings._element
    math_props = settings.find(qn("m:mathPr"))
    if math_props is None:
        math_props = OxmlElement("m:mathPr")
        settings.append(math_props)
    math_font = math_props.find(qn("m:mathFont"))
    if math_font is None:
        math_font = OxmlElement("m:mathFont")
        math_props.append(math_font)
    math_font.set(qn("m:val"), MATH_FONT)
    add_page_number(section)
    return doc


ALLOWED_BINARY = {
    ast.Add: lambda a, b: a + b,
    ast.Sub: lambda a, b: a - b,
    ast.Mult: lambda a, b: a * b,
    ast.Div: lambda a, b: a / b,
    ast.Pow: lambda a, b: a ** b,
}
ALLOWED_UNARY = {ast.UAdd: lambda a: a, ast.USub: lambda a: -a}


def safe_expression(expression):
    tree = ast.parse(expression, mode="eval")

    def evaluate(node, x):
        if isinstance(node, ast.Expression):
            return evaluate(node.body, x)
        if isinstance(node, ast.Name) and node.id == "x":
            return x
        if isinstance(node, ast.Constant) and isinstance(node.value, (int, float)):
            return float(node.value)
        if isinstance(node, ast.BinOp) and type(node.op) in ALLOWED_BINARY:
            return ALLOWED_BINARY[type(node.op)](evaluate(node.left, x), evaluate(node.right, x))
        if isinstance(node, ast.UnaryOp) and type(node.op) in ALLOWED_UNARY:
            return ALLOWED_UNARY[type(node.op)](evaluate(node.operand, x))
        if isinstance(node, ast.Call) and isinstance(node.func, ast.Name):
            if node.func.id == "sqrt" and len(node.args) == 1:
                return np.sqrt(evaluate(node.args[0], x))
            if node.func.id == "abs" and len(node.args) == 1:
                return np.abs(evaluate(node.args[0], x))
        raise ValueError("Graph expression contains an unsupported operation")

    return lambda x: evaluate(tree, x)


def font_path(bold=False):
    candidates = [
        "/usr/share/fonts/truetype/liberation2/LiberationSerif-Bold.ttf" if bold else
        "/usr/share/fonts/truetype/liberation2/LiberationSerif-Regular.ttf",
        "/System/Library/Fonts/Supplemental/Times New Roman Bold.ttf" if bold else
        "/System/Library/Fonts/Supplemental/Times New Roman.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return candidate
    raise RuntimeError("A Times-compatible graph font is unavailable")


def draw_graph_panel(draw, box, choice, xlim, ylim, asymptotes):
    x0, y0, x1, y1 = box
    pad_l, pad_r, pad_t, pad_b = 34, 20, 15, 38
    px0, py0, px1, py1 = x0 + pad_l, y0 + pad_t, x1 - pad_r, y1 - pad_b

    def tx(value):
        return px0 + (value - xlim[0]) / (xlim[1] - xlim[0]) * (px1 - px0)

    def ty(value):
        return py1 - (value - ylim[0]) / (ylim[1] - ylim[0]) * (py1 - py0)

    axis = (20, 20, 20)
    x_axis = min(max(ty(0), py0), py1)
    y_axis = min(max(tx(0), px0), px1)
    draw.line((px0, x_axis, px1, x_axis), fill=axis, width=1)
    draw.polygon([(px1, x_axis), (px1 - 9, x_axis - 4), (px1 - 9, x_axis + 4)], fill=axis)
    draw.line((y_axis, py1, y_axis, py0), fill=axis, width=1)
    draw.polygon([(y_axis, py0), (y_axis - 4, py0 + 9), (y_axis + 4, py0 + 9)], fill=axis)
    small = ImageFont.truetype(font_path(False), 22)
    bold = ImageFont.truetype(font_path(True), 34)
    draw.text((px1 - 2, x_axis + 3), "x", font=small, fill=axis, anchor="ra")
    draw.text((y_axis - 7, py0 - 2), "y", font=small, fill=axis, anchor="rt")
    for value in asymptotes:
        if xlim[0] <= value <= xlim[1]:
            xx = tx(value)
            yy = py0
            while yy < py1:
                draw.line((xx, yy, xx, min(yy + 8, py1)), fill=(100, 100, 100), width=1)
                yy += 16
    fn = safe_expression(choice["expression"])
    breakpoints = [xlim[0]] + sorted(v for v in asymptotes if xlim[0] < v < xlim[1]) + [xlim[1]]
    with np.errstate(all="ignore"):
        for start, end in zip(breakpoints, breakpoints[1:]):
            margin = max((xlim[1] - xlim[0]) * 0.001, 0.001)
            xs = np.linspace(start + margin, end - margin, 1200)
            ys = fn(xs)
            if np.isscalar(ys):
                ys = np.full_like(xs, ys)
            points = []
            for x_value, y_value in zip(xs, ys):
                if np.isfinite(y_value) and ylim[0] <= y_value <= ylim[1]:
                    point = (tx(float(x_value)), ty(float(y_value)))
                    if points and abs(point[1] - points[-1][1]) > 80:
                        if len(points) > 1:
                            draw.line(points, fill=(5, 5, 5), width=2)
                        points = []
                    points.append(point)
                else:
                    if len(points) > 1:
                        draw.line(points, fill=(5, 5, 5), width=2)
                    points = []
            if len(points) > 1:
                draw.line(points, fill=(5, 5, 5), width=2)
    draw.text(((x0 + x1) / 2, y1 - 5), choice["label"], font=bold,
              fill=(0, 0, 0), anchor="ms")


def create_graph_choices(spec, path):
    choices = spec.get("choices", [])
    labels = {choice.get("label") for choice in choices}
    if len(choices) != 4 or labels != {"I", "II", "III", "IV"}:
        raise ValueError("graph_choices must contain I, II, III, and IV")
    by_label = {choice["label"]: choice for choice in choices}
    xlim = [float(value) for value in spec.get("xlim", [-5, 5])]
    ylim = [float(value) for value in spec.get("ylim", [-5, 5])]
    asymptotes = [float(value) for value in spec.get("vertical_asymptotes", [])]
    image = Image.new("RGB", (1450, 900), "white")
    draw = ImageDraw.Draw(image)
    boxes = {
        "II": (35, 20, 710, 425),
        "I": (740, 20, 1415, 425),
        "IV": (35, 455, 710, 870),
        "III": (740, 455, 1415, 870),
    }
    for label, box in boxes.items():
        draw_graph_panel(draw, box, by_label[label], xlim, ylim, asymptotes)
    image.save(path, dpi=(200, 200))


def render_docx(question, docx_path, work_dir):
    doc = new_document()
    number = int(question.get("question_number", 6))
    if not 1 <= number <= 20:
        raise ValueError("question_number is out of range")
    if question.get("formula"):
        add_intro(doc, number, question.get("formula"))
    else:
        add_general_intro(doc, number, question.get("intro_parts", []))
    elements = question.get("elements", [])
    if not 4 <= len(elements) <= 20:
        raise ValueError("A worksheet must contain 4-20 layout elements")
    graph_index = 0
    for element in elements:
        element_type = element.get("type")
        after = min(max(int(element.get("after", 3)), 0), 12)
        if element_type == "paragraph":
            add_mixed(doc, element.get("parts", []), after=after,
                      bold=bool(element.get("bold", False)))
        elif element_type == "item":
            add_item(doc, str(element.get("label", "")), element.get("parts", []), after=after)
        elif element_type == "data_table":
            add_data_table(doc, element.get("headers", []), element.get("rows", []), after=after)
        elif element_type == "graph_choices":
            graph_index += 1
            graph_path = work_dir / f"graph-{graph_index}.png"
            create_graph_choices(element, graph_path)
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(after)
            paragraph.add_run().add_picture(str(graph_path), width=Cm(14.0))
        else:
            raise ValueError(f"Unsupported layout element: {element_type}")
    doc.save(docx_path)


def fontconfig_xml(cache_dir):
    dirs = [FONT_DIR, Path("/usr/share/fonts"), Path("/usr/local/share/fonts")]
    body = "\n".join(f"  <dir>{path}</dir>" for path in dirs if path.exists())
    return (
        '<?xml version="1.0"?>\n'
        '<!DOCTYPE fontconfig SYSTEM "fonts.dtd">\n'
        '<fontconfig>\n'
        f'{body}\n'
        f'  <cachedir>{cache_dir}</cachedir>\n'
        '</fontconfig>\n'
    )


def convert_and_verify(docx_path, pdf_path, preview_dir, temp_dir,
                       document_type="summary_question"):
    cache = temp_dir / "font-cache"
    cache.mkdir()
    config = temp_dir / "fonts.conf"
    config.write_text(fontconfig_xml(cache), encoding="utf-8")
    converted = temp_dir / "converted"
    converted.mkdir()
    profile = temp_dir / "lo-profile"
    env = os.environ.copy()
    env["FONTCONFIG_FILE"] = str(config)
    env["XDG_CACHE_HOME"] = str(cache)
    subprocess.run([
        executable("soffice"), f"-env:UserInstallation=file://{profile}",
        "--headless", "--convert-to", "pdf", "--outdir", str(converted),
        str(docx_path),
    ], check=True, env=env, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    generated = converted / f"{docx_path.stem}.pdf"
    if not generated.exists():
        raise RuntimeError("LibreOffice did not produce a PDF")
    shutil.copy2(generated, pdf_path)
    preview_dir.mkdir(parents=True, exist_ok=True)
    subprocess.run([
        executable("pdftoppm"), "-png", "-r", "150", str(pdf_path),
        str(preview_dir / "page"),
    ], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    info = subprocess.run([executable("pdfinfo"), str(pdf_path)], check=True,
                          text=True, stdout=subprocess.PIPE).stdout
    page_match = re.search(r"^Pages:\s+(\d+)$", info, flags=re.MULTILINE)
    page_count = int(page_match.group(1)) if page_match else 0
    max_pages = 3 if document_type == "practice_worksheet" else 1
    if not 1 <= page_count <= max_pages or "Page size:       595" not in info:
        raise RuntimeError(
            f"Generated {document_type} must contain 1-{max_pages} A4 pages"
        )
    fonts = subprocess.run([executable("pdffonts"), str(pdf_path)], check=True,
                           text=True, stdout=subprocess.PIPE).stdout
    if "FbDavidNewPro-Regular" not in fonts or "FbDavidNewPro-Bold" not in fonts:
        raise RuntimeError("Exact David New Pro fonts were not embedded")


def safe_slug(value):
    value = re.sub(r"[^a-zA-Z0-9-]+", "-", str(value)).strip("-")
    return value[:60] or "question"


def main():
    if len(sys.argv) != 3:
        raise SystemExit("Usage: render_question.py SPEC_JSON OUTPUT_DIR")
    spec_path = Path(sys.argv[1]).resolve()
    output_dir = Path(sys.argv[2]).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    payload = json.loads(spec_path.read_text(encoding="utf-8"))
    questions = payload.get("questions", [])
    if not 1 <= len(questions) <= 3:
        raise ValueError("Expected 1-3 questions")
    results = []
    for index, question in enumerate(questions, 1):
        token = safe_slug(question.get("id", f"question-{index}"))
        docx_path = output_dir / f"{token}.docx"
        pdf_path = output_dir / f"{token}.pdf"
        preview_dir = output_dir / f"{token}-preview"
        with tempfile.TemporaryDirectory(prefix="gradeup-471-") as temp_name:
            temp_dir = Path(temp_name)
            document_type = str(question.get("document_type", "summary_question"))
            if document_type == "practice_worksheet":
                render_practice_worksheet_docx(question, docx_path)
            elif document_type == "summary_question":
                render_docx(question, docx_path, temp_dir)
            else:
                raise ValueError("Unsupported document_type")
            convert_and_verify(
                docx_path, pdf_path, preview_dir, temp_dir,
                document_type=document_type,
            )
        preview_pages = sorted(preview_dir.glob("page-*.png"))
        results.append({
            "id": token,
            "pdf": pdf_path.name,
            "preview": f"{preview_dir.name}/{preview_pages[0].name}",
            "title": str(question.get("title", "שאלת חקירה 471")),
        })
    print(json.dumps({"success": True, "files": results}, ensure_ascii=False))


if __name__ == "__main__":
    main()
